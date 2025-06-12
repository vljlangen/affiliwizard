import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
import re
from docx.oxml.ns import qn
from openpyxl import Workbook

# Superscript mapping for 1-9
SUPERSCRIPTS = {
    '0': '\u2070', '1': '\u00b9', '2': '\u00b2', '3': '\u00b3',
    '4': '\u2074', '5': '\u2075', '6': '\u2076', '7': '\u2077', '8': '\u2078', '9': '\u2079'
}

def to_superscript(number):
    return ''.join(SUPERSCRIPTS[d] for d in str(number))

def to_superscript_seq(seq):
    # Converts a string like '1,2' to '¹,²'
    SUPERSCRIPTS = {
        '0': '\u2070', '1': '\u00b9', '2': '\u00b2', '3': '\u00b3',
        '4': '\u2074', '5': '\u2075', '6': '\u2076', '7': '\u2077', '8': '\u2078', '9': '\u2079', ',': ','
    }
    return ''.join(SUPERSCRIPTS.get(ch, ch) for ch in seq)

def parse_data(raw_text):
    lines = [line.strip() for line in raw_text.strip().split('\n') if line.strip()]
    authors = []
    affil_map = {}
    affil_list = []
    affil_counter = 1
    author_affils = []
    for line in lines:
        cols = re.split(r'\t+', line)
        if not cols or not cols[0]:
            continue
        author = cols[0].strip()
        affils = [a.strip() for a in cols[1:] if a.strip()]
        affil_nums = []
        for affil in affils:
            if affil not in affil_map:
                affil_map[affil] = affil_counter
                affil_list.append(affil)
                affil_counter += 1
            affil_nums.append(affil_map[affil])
        author_affils.append((author, affil_nums))
    return author_affils, affil_list, affil_map

def format_block(author_affils, affil_list, affil_map, plain_text=False):
    # Author line with all superscript numbers and commas (default)
    author_line = []
    for author, nums in author_affils:
        sorted_nums = sorted(nums)
        num_seq = ','.join(str(n) for n in sorted_nums)
        if plain_text:
            # Use (1,2) format for plain text
            author_line.append(f"{author} ({num_seq})" if num_seq else author)
        else:
            supers = to_superscript_seq(num_seq)
            author_line.append(f"{author}{supers}")
    author_line_str = ', '.join(author_line)
    # Affiliation list
    if plain_text:
        affil_lines = [f"{affil_map[affil]}. {affil}" for affil in affil_list]
    else:
        affil_lines = [f"{affil_map[affil]} {affil}" for affil in affil_list]
    return f"{author_line_str}\n\n" + '\n'.join(affil_lines)

def create_new_docx(docx_path, author_affils, affil_list, affil_map):
    doc = Document()
    # Set default font to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # Author line with true superscript for numbers/commas
    p = doc.add_paragraph()
    for idx, (author, nums) in enumerate(author_affils):
        sorted_nums = sorted(nums)
        run = p.add_run(author)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if sorted_nums:
            num_seq = ','.join(str(n) for n in sorted_nums)
            sup_run = p.add_run(num_seq)
            sup_run.font.superscript = True
            sup_run.font.name = 'Arial'
            sup_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if idx < len(author_affils) - 1:
            comma_run = p.add_run(", ")
            comma_run.font.name = 'Arial'
            comma_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    doc.add_paragraph("")  # Blank line
    # Affiliation list with numbers and period after digit
    for affil in affil_list:
        para = doc.add_paragraph(f"{affil_map[affil]}. {affil}")
        for run in para.runs:
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    doc.save(docx_path)

def create_new_html(html_path, author_affils, affil_list, affil_map):
    def html_superscript(seq):
        # Converts a string like '1,2' to HTML superscript: <sup>1,2</sup>
        return f"<sup>{seq}</sup>" if seq else ""

    # Author line with superscript numbers/commas
    author_line = []
    for author, nums in author_affils:
        sorted_nums = sorted(nums)
        num_seq = ','.join(str(n) for n in sorted_nums)
        supers = html_superscript(num_seq)
        author_line.append(f"{author}{supers}")
    author_line_str = ', '.join(author_line)
    # Affiliation list with numbers and period after digit
    affil_lines = [f"{affil_map[affil]}. {affil}" for affil in affil_list]
    affil_html = '<br>\n'.join(affil_lines)
    html = f"""<html>
<head>
<meta charset='utf-8'>
<title>Author Affiliations</title>
<style>
body {{ font-family: Arial, sans-serif; }}
</style>
</head>
<body>
<p>{author_line_str}</p>
<br>
<p>{affil_html}</p>
</body>
</html>"""
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html)

def create_new_txt(txt_path, author_affils, affil_list, affil_map):
    text = format_block(author_affils, affil_list, affil_map, plain_text=True)
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text)

class AuthorAffiliationsApp:
    def __init__(self, root):
        self.root = root
        root.title("Author Affiliations File Creator")
        usage_text = (
            "How to use:\n\n"
            "1. Paste your author-affiliation data from Excel to the text box below.\n"
            "2. Then use the buttons below to export author-affiliation data to .docx, .html, or .txt\n\n"
            "Format instructions for the Excel file:\n\n"
            "- Each line: Author Name [tab] Affiliation 1 [tab] Affiliation 2 ...\n\n"
            "- Practical example:\n\n"
            "John Doe\tBogus Institute, CA, USA\tExample Institute, TS, USA\n\n\n\n"
            "Download an Excel example file below if needed.\n\n"
            "\nLicense:"
            "\nMIT License, Copyright (c) 2025 Ville Langén.\n"
            "(TLDR: MIT License means you can use this freely. Just credit me and don’t sue me.)"
        )
        self.usage_label = tk.Label(root, text=usage_text, justify="left", anchor="w", fg="blue")
        self.usage_label.pack(padx=10, pady=(10,0), anchor="w")
        self.text = scrolledtext.ScrolledText(root, width=80, height=20)
        self.text.pack(padx=10, pady=10)
        self.create_btn = tk.Button(root, text="Create .docx File", command=self.process)
        self.create_btn.pack(pady=10)
        self.create_html_btn = tk.Button(root, text="Create HTML File", command=self.process_html)
        self.create_html_btn.pack(pady=5)
        self.create_txt_btn = tk.Button(root, text="Create Plain Text File", command=self.process_txt)
        self.create_txt_btn.pack(pady=5)
        self.example_excel_btn = tk.Button(root, text="Download Example Excel File", command=self.download_example_excel)
        self.example_excel_btn.pack(pady=5)
        self.quit_btn = tk.Button(root, text="Quit", command=root.quit, fg="red")
        self.quit_btn.pack(pady=(20,10))

    def process(self):
        raw_text = self.text.get("1.0", tk.END)
        if not raw_text.strip():
            messagebox.showerror("Error", "Please paste author-affiliation data.")
            return
        try:
            author_affils, affil_list, affil_map = parse_data(raw_text)
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], title="Save new .docx file")
            if not save_path:
                return
            create_new_docx(save_path, author_affils, affil_list, affil_map)
            messagebox.showinfo("Success", f"New .docx file created at:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def process_html(self):
        raw_text = self.text.get("1.0", tk.END)
        if not raw_text.strip():
            messagebox.showerror("Error", "Please paste author-affiliation data.")
            return
        try:
            author_affils, affil_list, affil_map = parse_data(raw_text)
            save_path = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML Files", "*.html")], title="Save new HTML file")
            if not save_path:
                return
            create_new_html(save_path, author_affils, affil_list, affil_map)
            messagebox.showinfo("Success", f"New HTML file created at:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def process_txt(self):
        raw_text = self.text.get("1.0", tk.END)
        if not raw_text.strip():
            messagebox.showerror("Error", "Please paste author-affiliation data.")
            return
        try:
            author_affils, affil_list, affil_map = parse_data(raw_text)
            save_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")], title="Save new text file")
            if not save_path:
                return
            create_new_txt(save_path, author_affils, affil_list, affil_map)
            messagebox.showinfo("Success", f"New text file created at:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def download_example_excel(self):
        save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")], title="Save example Excel file")
        if not save_path:
            return
        try:
            create_example_excel(save_path)
            messagebox.showinfo("Success", f"Example Excel file saved at:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

# Helper to create an example Excel file
def create_example_excel(xlsx_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Authors"
    ws.append(["John Doe, MD", "Bogus Institute, CA, USA", "Example Institute, TS, USA"])
    ws.append(["Jane Doe, MD", "Sample University, TX, USA"])
    wb.save(xlsx_path)

def main():
    root = tk.Tk()
    app = AuthorAffiliationsApp(root)
    root.mainloop()

if __name__ == "__main__":
    main() 